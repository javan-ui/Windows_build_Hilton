#!/usr/bin/env python3
"""
Hilton High School — Report Card Splitter v3
Tabs: Generate | Pay Codes | Defaulters
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import sys
import csv
import json
import re
import random
import string
import difflib
import subprocess
import shutil
import tempfile
import time
import zipfile
import urllib.request
import urllib.error
from datetime import datetime
from pathlib import Path


# ═══════════════════════════════════════════════════════════════════════════════
# SILENT AUTO-INSTALL (Windows-safe — no console flash)
# ═══════════════════════════════════════════════════════════════════════════════

def _ensure(import_name: str, pip_name: str):
    try:
        __import__(import_name)
    except ImportError:
        try:
            kw: dict = {
                "stdout": subprocess.DEVNULL,
                "stderr": subprocess.DEVNULL,
            }
            if sys.platform == "win32":
                kw["creationflags"] = 0x08000000  # CREATE_NO_WINDOW
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install",
                 pip_name, "--quiet", "--no-warn-script-location"],
                **kw
            )
        except Exception:
            pass  # We'll surface a proper error when the import is attempted


_ensure("fitz",    "PyMuPDF")
_ensure("PIL",     "Pillow")
_ensure("openpyxl","openpyxl")
_ensure("docx",    "python-docx")

try:
    import fitz
    from PIL import Image, ImageTk
except ImportError as _exc:
    # Show a user-friendly dialog before giving up
    try:
        import tkinter as _tk, tkinter.messagebox as _mb
        _r = _tk.Tk(); _r.withdraw()
        _mb.showerror("Missing component",
            f"A required component could not be installed:\n{_exc}\n\n"
            "Please run 'Card Splitter Setup.bat' first, then try again.")
        _r.destroy()
    except Exception:
        pass
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════════
# THEME
# ═══════════════════════════════════════════════════════════════════════════════

PRIMARY      = "#1a6b3a"
PRIMARY_DARK = "#114428"
PRIMARY_PALE = "#e8f5ee"
ACCENT       = "#2563eb"
RED          = "#dc2626"
RED_PALE     = "#fef2f2"
AMBER        = "#b45309"
BG           = "#f1f5f9"
CARD         = "#ffffff"
BORDER       = "#e2e8f0"
TEXT         = "#1e293b"
TEXT_MUTED   = "#64748b"
BTN_BG       = "#e8edf3"

FONT      = ("Segoe UI", 9)
FONT_BOLD = ("Segoe UI", 9,  "bold")
FONT_LG   = ("Segoe UI", 11, "bold")
FONT_SM   = ("Segoe UI", 8)

MAX_REPORT_STORAGE = 5 * 1024 * 1024 * 1024
STORAGE_WARNING_BYTES = int(MAX_REPORT_STORAGE * 0.8)
DEFAULT_UPLOAD_PASSWORD = "Hilt0n.High.o3"
DEFAULT_NETLIFY_TOKEN = "nfp_E4ZR9rQEy9Wc2szLqcvff2gWPNXmCWcqf93a"
DEFAULT_NETLIFY_SITE_ID = "bdfc8f16-e305-46ce-a0eb-87186a403e8d"


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def random_number(length=6):
    return "".join(random.choices(string.digits, k=length))


def safe_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]', "", name)
    return re.sub(r"\s+", " ", name).strip() or "Unknown"


def normalize_name(name: str) -> str:
    return re.sub(r"\s+", " ", name.lower().strip())


def is_numeric_code(val: str) -> bool:
    return bool(re.match(r"^\d{4,}$", val.strip()))


def fuzzy_match(target: str, candidates, threshold=0.78):
    norm  = normalize_name(target)
    cands = list(candidates)
    if norm in cands:
        return norm
    norm_sorted = " ".join(sorted(norm.split()))
    for key in cands:
        if " ".join(sorted(key.split())) == norm_sorted:
            return key
    matches = difflib.get_close_matches(norm, cands, n=1, cutoff=threshold)
    return matches[0] if matches else None


# ═══════════════════════════════════════════════════════════════════════════════
# FILE PARSERS
# ═══════════════════════════════════════════════════════════════════════════════

def _rows_to_pairs(rows):
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    if n_cols < 2:
        return []

    def num_score(col):
        vals = [r[col].strip() for r in rows if col < len(r) and r[col].strip()]
        return (sum(1 for v in vals if is_numeric_code(v)) / len(vals)) if vals else 0.0

    scores   = [num_score(i) for i in range(n_cols)]
    code_col = max(range(n_cols), key=lambda i: scores[i])
    if scores[code_col] < 0.4:
        return []

    name_col, best = None, -1
    for i in range(n_cols):
        if i == code_col:
            continue
        vals = [r[i].strip() for r in rows if i < len(r) and r[i].strip()]
        if not vals:
            continue
        ts = sum(1 for v in vals if not is_numeric_code(v) and len(v) > 2) / len(vals)
        if ts > best:
            best, name_col = ts, i

    if name_col is None:
        return []

    start = 0
    if code_col < len(rows[0]) and not is_numeric_code(rows[0][code_col].strip()):
        start = 1

    pairs = []
    for row in rows[start:]:
        name = row[name_col].strip() if name_col < len(row) else ""
        code = row[code_col].strip() if code_col < len(row) else ""
        if name and code and is_numeric_code(code):
            pairs.append((name, code))
    return pairs


def parse_paycode_file(path: str):
    ext = Path(path).suffix.lower()
    try:
        if ext == ".csv":
            return _parse_csv_codes(path)
        elif ext in (".xlsx", ".xls"):
            return _parse_excel_codes(path)
        elif ext in (".docx", ".doc"):
            return _parse_docx_codes(path)
        elif ext == ".pdf":
            return _parse_pdf_codes(path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    except (RuntimeError, ValueError):
        raise
    except Exception as exc:
        raise RuntimeError(f"Error reading {Path(path).name}: {exc}")


def _parse_csv_codes(path):
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(path, newline="", encoding=enc) as f:
                return _rows_to_pairs([r for r in csv.reader(f)])
        except UnicodeDecodeError:
            continue
    return []


def _parse_excel_codes(path):
    import openpyxl
    wb   = openpyxl.load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
        break
    wb.close()
    return _rows_to_pairs(rows)


def _parse_docx_codes(path):
    import docx as _docx
    doc  = _docx.Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
    if not rows:
        for para in doc.paragraphs:
            text = para.text.strip()
            if "\t" in text:
                rows.append(text.split("\t"))
            elif "," in text:
                rows.append(text.split(","))
    return _rows_to_pairs(rows)


def _parse_pdf_codes(path):
    doc  = fitz.open(path)
    rows = []
    for page in doc:
        words = page.get_text("words")
        if not words:
            continue
        lines: dict = {}
        for w in words:
            y_key = round(w[1] / 5) * 5
            lines.setdefault(y_key, []).append(w)
        for y in sorted(lines):
            lw  = sorted(lines[y], key=lambda w: w[0])
            row, cur, prev_x1 = [], [], None
            for w in lw:
                if prev_x1 is not None and w[0] - prev_x1 > 20:
                    row.append(" ".join(cur)); cur = []
                cur.append(w[4]); prev_x1 = w[2]
            if cur:
                row.append(" ".join(cur))
            if len(row) >= 2:
                rows.append(row)
    doc.close()
    return _rows_to_pairs(rows)


def parse_names_file(path: str):
    ext   = Path(path).suffix.lower()
    names = []
    skip  = {"name", "student", "student name", "no", "no.", "sn", "#", "sr"}
    try:
        if ext == ".csv":
            for enc in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    with open(path, newline="", encoding=enc) as f:
                        for row in csv.reader(f):
                            for cell in row:
                                c = cell.strip()
                                if c and not is_numeric_code(c):
                                    names.append(c)
                    break
                except UnicodeDecodeError:
                    continue
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        val = str(cell).strip() if cell is not None else ""
                        if val and val.lower() != "none" and not is_numeric_code(val):
                            names.append(val)
                break
            wb.close()
        elif ext in (".docx", ".doc"):
            import docx as _docx
            doc = _docx.Document(path)
            for para in doc.paragraphs:
                t = para.text.strip()
                if t and not is_numeric_code(t):
                    names.append(t)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and not is_numeric_code(t):
                            names.append(t)
        elif ext == ".pdf":
            doc = fitz.open(path)
            for page in doc:
                for line in page.get_text("text").split("\n"):
                    line = line.strip()
                    if line and not is_numeric_code(line) and len(line) > 2:
                        names.append(line)
            doc.close()
    except (RuntimeError, ValueError):
        raise
    except Exception as exc:
        raise RuntimeError(f"Error reading {Path(path).name}: {exc}")

    return [n for n in names if normalize_name(n) not in skip and len(n) >= 3]


# ═══════════════════════════════════════════════════════════════════════════════
# DATA STORE
# ═══════════════════════════════════════════════════════════════════════════════

class DataStore:
    APP_DIR = Path(os.environ.get("APPDATA", Path.home())) / "HiltonHighCardSplitter"

    def __init__(self):
        self.APP_DIR.mkdir(parents=True, exist_ok=True)
        self._codes_path    = self.APP_DIR / "paycodes.json"
        self._default_path  = self.APP_DIR / "defaulters.json"
        self._settings_path = self.APP_DIR / "settings.json"
        self._history_path  = self.APP_DIR / "report_history.json"
        self.report_archive = self.APP_DIR / "report_archive"
        self.report_archive.mkdir(parents=True, exist_ok=True)
        self._load()

    def _load(self):
        # Pay codes
        if self._codes_path.exists():
            try:
                d = json.loads(self._codes_path.read_text("utf-8"))
                self.students = d.get("students", {})
            except Exception:
                self.students = {}
        else:
            self.students = {}

        # Defaulters
        if self._default_path.exists():
            try:
                d = json.loads(self._default_path.read_text("utf-8"))
                self._defaulters        = set(d.get("defaulters", []))
                self._defaulter_display = d.get("display", {})
            except Exception:
                self._defaulters, self._defaulter_display = set(), {}
        else:
            self._defaulters, self._defaulter_display = set(), {}

        try:
            self.settings = json.loads(self._settings_path.read_text("utf-8"))
        except Exception:
            self.settings = {}
        self.settings.setdefault("upload_password", DEFAULT_UPLOAD_PASSWORD)
        self.settings.setdefault("netlify_token", DEFAULT_NETLIFY_TOKEN)
        self.settings.setdefault("netlify_site_id", DEFAULT_NETLIFY_SITE_ID)
        self.settings.setdefault("copyright_year", str(datetime.now().year))

        try:
            self.report_history = json.loads(self._history_path.read_text("utf-8"))
            if not isinstance(self.report_history, list):
                self.report_history = []
        except Exception:
            self.report_history = []

    def _save_settings(self):
        self._settings_path.write_text(
            json.dumps(self.settings, indent=2, ensure_ascii=False), "utf-8")

    def _save_history(self):
        self._history_path.write_text(
            json.dumps(self.report_history, indent=2, ensure_ascii=False), "utf-8")

    def storage_bytes(self) -> int:
        total = 0
        for root, _, files in os.walk(self.report_archive):
            for name in files:
                try:
                    total += (Path(root) / name).stat().st_size
                except OSError:
                    pass
        return total

    @staticmethod
    def human_size(size: int) -> str:
        value = float(size)
        for unit in ("B", "KB", "MB", "GB"):
            if value < 1024 or unit == "GB":
                return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
            value /= 1024
        return f"{value:.1f} GB"

    def add_report_batch(self, title: str, source_dir: str, filenames, manifest_rows=None) -> dict:
        files = [Path(source_dir) / name for name in filenames]
        batch_bytes = sum(p.stat().st_size for p in files if p.exists())
        current = self.storage_bytes()
        if current + batch_bytes > MAX_REPORT_STORAGE:
            raise RuntimeError(
                "The saved report limit of 5 GB would be exceeded. "
                "Delete an older report batch before saving this one.")
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        batch_dir = self.report_archive / f"{stamp}_{safe_filename(title)}"
        batch_dir.mkdir(parents=True, exist_ok=True)
        copied = []
        for source in files:
            if source.exists():
                target = batch_dir / source.name
                shutil.copy2(source, target)
                copied.append(target.name)
        record = {
            "title": title,
            "date": datetime.now().isoformat(timespec="seconds"),
            "size": batch_bytes,
            "folder": str(batch_dir),
            "files": copied,
            "manifest_rows": manifest_rows or [],
        }
        self.report_history.insert(0, record)
        self._save_history()
        return record

    def delete_report_batch(self, record: dict):
        folder = Path(record.get("folder", ""))
        if folder.exists() and folder.is_dir():
            shutil.rmtree(folder, ignore_errors=True)
        self.report_history = [
            item for item in self.report_history
            if item.get("folder") != record.get("folder")]
        self._save_history()

    def save_settings(self, values: dict):
        for key in ("upload_password", "netlify_token", "netlify_site_id", "copyright_year"):
            if values.get(key) is not None:
                self.settings[key] = str(values[key]).strip()
        self._save_settings()

    def export_archive(self, destination: str):
        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in self.APP_DIR.rglob("*"):
                if path.is_file() and path.resolve() != Path(destination).resolve():
                    archive.write(path, Path("app_data") / path.relative_to(self.APP_DIR))
            for index, record in enumerate(self.report_history):
                folder = Path(record.get("folder", ""))
                if folder.exists():
                    for path in folder.rglob("*"):
                        if path.is_file():
                            archive.write(path, Path("reports") / str(index) / path.name)

    def import_archive(self, source: str):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            with zipfile.ZipFile(source) as archive:
                for member in archive.infolist():
                    target = (root / member.filename).resolve()
                    if not str(target).startswith(str(root.resolve())):
                        raise RuntimeError("The import archive contains an unsafe path.")
                    if not member.is_dir():
                        target.parent.mkdir(parents=True, exist_ok=True)
                        with archive.open(member) as src, open(target, "wb") as dst:
                            shutil.copyfileobj(src, dst)
            imported = root / "app_data"
            if imported.exists():
                for path in imported.rglob("*"):
                    if path.is_file():
                        target = self.APP_DIR / path.relative_to(imported)
                        target.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(path, target)
        self._load()


    def _save_codes(self):
        self._codes_path.write_text(
            json.dumps({"students": self.students}, indent=2, ensure_ascii=False), "utf-8")

    def _save_defaulters(self):
        self._default_path.write_text(
            json.dumps({"defaulters": list(self._defaulters),
                        "display":    self._defaulter_display},
                       indent=2, ensure_ascii=False), "utf-8")

    # ── Pay codes ──────────────────────────────────────────────────────────────

    def import_pairs(self, pairs) -> int:
        added = 0
        for name, code in pairs:
            existing_key = fuzzy_match(name, self.students)
            key = existing_key if existing_key else normalize_name(name)
            if key not in self.students:
                self.students[key] = {"display": name, "codes": []}
            if code not in self.students[key]["codes"]:
                self.students[key]["codes"].append(code)
                added += 1
        self._save_codes()
        return added

    def add_code(self, display_name: str, code: str):
        key = normalize_name(display_name)
        if key not in self.students:
            self.students[key] = {"display": display_name, "codes": []}
        if code not in self.students[key]["codes"]:
            self.students[key]["codes"].append(code)
        self._save_codes()

    def get_codes(self, name: str):
        key = fuzzy_match(name, self.students)
        return list(self.students[key]["codes"]) if key else []

    def all_students(self):
        return sorted(
            [(v["display"], v["codes"]) for v in self.students.values()],
            key=lambda x: x[0].lower())

    def remove_student(self, display_name: str):
        key = normalize_name(display_name)
        if key in self.students:
            del self.students[key]
            self._save_codes()

    def clear_all_codes(self):
        self.students = {}
        self._save_codes()

    # ── Defaulters ─────────────────────────────────────────────────────────────

    def add_defaulter(self, name: str):
        key = normalize_name(name)
        self._defaulters.add(key)
        self._defaulter_display[key] = name
        self._save_defaulters()

    def import_defaulters(self, names):
        for name in names:
            self.add_defaulter(name)

    def remove_defaulter(self, name: str):
        key = normalize_name(name)
        self._defaulters.discard(key)
        self._defaulter_display.pop(key, None)
        self._save_defaulters()

    def is_defaulter(self, name: str) -> bool:
        return bool(fuzzy_match(name, self._defaulters))

    def all_defaulters(self):
        return sorted(
            [self._defaulter_display.get(k, k) for k in self._defaulters],
            key=str.lower)


# ═══════════════════════════════════════════════════════════════════════════════
# MISSING CODES DIALOG
# ═══════════════════════════════════════════════════════════════════════════════

class MissingCodesDialog(tk.Toplevel):
    def __init__(self, parent, missing_names, store: DataStore):
        super().__init__(parent)
        self.title("Missing School Pay Codes")
        self.geometry("580x520")
        self.minsize(480, 400)
        self.resizable(True, True)
        self.grab_set()
        self.configure(bg=BG)
        self._store     = store
        self._missing   = list(missing_names)
        self._code_vars: dict = {}
        self.result     = "cancel"
        self._mode      = tk.StringVar(value="manual")
        self._build()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg="#92400e", pady=10)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Missing School Pay Codes",
                 bg="#92400e", fg="white", font=FONT_LG).pack(padx=14, anchor="w")
        n = len(self._missing)
        tk.Label(hdr,
                 text=f"{n} student{'s' if n != 1 else ''} have no pay code in stored data.",
                 bg="#92400e", fg="#fef3c7", font=FONT).pack(padx=14, anchor="w")

        # Mode radio buttons
        rr = tk.Frame(self, bg=BG, pady=8)
        rr.pack(fill="x", padx=14)
        tk.Label(rr, text="What would you like to do?",
                 bg=BG, font=FONT_BOLD, fg=TEXT).pack(anchor="w", pady=(0, 4))
        for val, label in [
            ("manual", "Enter codes manually"),
            ("upload", "Upload a file with the missing codes"),
            ("skip",   "Skip these students (no report cards for them)"),
        ]:
            tk.Radiobutton(rr, text=label, variable=self._mode, value=val,
                           command=self._switch_mode,
                           bg=BG, font=FONT, activebackground=BG,
                           cursor="hand2").pack(anchor="w", padx=16)

        self._content = tk.Frame(self, bg=BG)
        self._content.pack(fill="both", expand=True, padx=14, pady=(4, 4))

        # Buttons
        bar = tk.Frame(self, bg=BORDER, pady=8)
        bar.pack(fill="x", side="bottom")
        _btn(bar, "Cancel", self._on_cancel, bg=BTN_BG).pack(side="right", padx=(0, 14))
        _btn(bar, "Continue  >", self._on_confirm,
             bg=PRIMARY, fg="white", font=FONT_BOLD).pack(side="right", padx=(0, 6))

        self._switch_mode()

    def _clear(self):
        for w in self._content.winfo_children():
            w.destroy()

    def _switch_mode(self):
        self._clear()
        mode = self._mode.get()
        if mode == "manual":
            self._build_manual()
        elif mode == "upload":
            self._build_upload()
        else:
            self._build_skip()

    def _build_manual(self):
        tk.Label(self._content, text="Enter a pay code next to each student name:",
                 bg=BG, font=FONT, fg=TEXT_MUTED).pack(anchor="w", pady=(4, 6))
        outer = tk.Frame(self._content, bg=BG)
        outer.pack(fill="both", expand=True)
        cv  = tk.Canvas(outer, bg=BG, highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient="vertical", command=cv.yview)
        cv.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        cv.pack(fill="both", expand=True)
        inner = tk.Frame(cv, bg=BG)
        cv.create_window((0, 0), window=inner, anchor="nw")
        self._code_vars = {}
        for name in self._missing:
            row = tk.Frame(inner, bg=BG)
            row.pack(fill="x", pady=3)
            tk.Label(row, text=name, bg=BG, font=FONT_BOLD, fg=TEXT,
                     width=30, anchor="w").pack(side="left")
            var = tk.StringVar()
            self._code_vars[name] = var
            tk.Entry(row, textvariable=var, font=FONT,
                     bd=1, relief="solid", width=14, bg="white").pack(side="left", padx=(8, 0))
        inner.update_idletasks()
        cv.configure(scrollregion=cv.bbox("all"))

    def _build_upload(self):
        tk.Label(self._content,
                 text="Choose a CSV, Excel, Word or PDF file containing\n"
                      "student names and their school pay codes.",
                 bg=BG, font=FONT, fg=TEXT_MUTED, justify="left").pack(anchor="w", pady=(8, 10))
        self._upload_status = tk.Label(self._content, text="No file selected.",
                                       bg=BG, font=FONT_SM, fg=TEXT_MUTED,
                                       wraplength=440, justify="left")
        self._upload_status.pack(anchor="w", pady=(0, 8))
        _btn(self._content, "Browse for File", self._do_upload,
             bg=ACCENT, fg="white").pack(anchor="w")

    def _do_upload(self):
        path = filedialog.askopenfilename(
            parent=self, title="Select file with school pay codes",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("All", "*.*")])
        if not path:
            return
        try:
            pairs = parse_paycode_file(path)
            added = self._store.import_pairs(pairs)
            self._upload_status.config(
                text=f"OK  {Path(path).name} — {len(pairs)} pairs read, {added} new codes stored.",
                fg=PRIMARY)
        except Exception as exc:
            self._upload_status.config(text=f"Error: {exc}", fg=RED)

    def _build_skip(self):
        n = len(self._missing)
        tk.Label(self._content,
                 text=f"These {n} student{'s' if n != 1 else ''} will be skipped:\n"
                      "No report card will be generated for them.",
                 bg=BG, font=FONT, fg=TEXT_MUTED, justify="left").pack(anchor="w", pady=(8, 8))
        frm = tk.Frame(self._content, bg=BG)
        frm.pack(fill="both", expand=True)
        sb = tk.Scrollbar(frm); sb.pack(side="right", fill="y")
        lb = tk.Listbox(frm, font=FONT, bg=CARD, bd=1, relief="solid",
                        yscrollcommand=sb.set, highlightthickness=0)
        lb.pack(fill="both", expand=True)
        sb.config(command=lb.yview)
        for name in self._missing:
            lb.insert("end", name)

    def _on_confirm(self):
        mode = self._mode.get()
        if mode == "manual":
            for name, var in self._code_vars.items():
                code = var.get().strip()
                if code:
                    self._store.add_code(name, code)
        self.result = "skip" if mode == "skip" else "continue"
        self.destroy()

    def _on_cancel(self):
        self.result = "cancel"
        self.destroy()


# ═══════════════════════════════════════════════════════════════════════════════
# STANDALONE BUTTON HELPER (outside class, used in dialogs)
# ═══════════════════════════════════════════════════════════════════════════════

def _btn(parent, text, command, bg=BTN_BG, fg=TEXT, font=None, padx=10, pady=5):
    return tk.Button(parent, text=text, command=command, bg=bg, fg=fg,
                     font=font or FONT, relief="flat", padx=padx, pady=pady,
                     cursor="hand2", activebackground=bg, activeforeground=fg)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN APPLICATION
# ═══════════════════════════════════════════════════════════════════════════════

class ReportSplitter(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Hilton High School - Report Card Splitter")
        self.geometry("1380x830")
        self.minsize(1040, 680)
        self.configure(bg=BG)

        self._store = DataStore()

        # Generate-tab state
        self.pdf_doc         = None
        self.current_path    = ""
        self.current_page    = 0
        self.page_scale      = 1.0
        self.page_photo      = None
        self.sel_start       = None
        self.sel_rect_id     = None
        self.name_rects:dict = {}
        self.same_region_var = tk.BooleanVar(value=True)
        self.circular_doc    = None
        self.circular_path   = ""
        self.medical_path    = ""
        self.banner_path     = ""
        self.processing      = False
        self._secret_key_presses = []
        self._upload_unlocked = False
        self._settings_tab = None

        # Per-file report card titles
        self.file_report_names: dict = {}

        self._apply_theme()
        self._build_ui()
        self.after(100, self._fit_canvas)

    # ── Theme ──────────────────────────────────────────────────────────────────

    def _apply_theme(self):
        try:
            s = ttk.Style(self)
            try:
                s.theme_use("clam")
            except Exception:
                try:
                    s.theme_use("alt")
                except Exception:
                    pass
            s.configure("TNotebook",
                        background=PRIMARY_DARK, borderwidth=0,
                        tabmargins=[0, 0, 0, 0])
            s.configure("TNotebook.Tab",
                        background=PRIMARY_DARK, foreground="#8fc8a8",
                        padding=[20, 10], font=("Segoe UI", 10, "bold"),
                        borderwidth=0)
            s.map("TNotebook.Tab",
                  background=[("selected", PRIMARY), ("active", PRIMARY)],
                  foreground=[("selected", "white"), ("active", "white")])
            s.configure("TProgressbar",
                        troughcolor=BORDER, background=PRIMARY,
                        borderwidth=0, thickness=8)
            s.configure("Treeview",
                        background=CARD, fieldbackground=CARD,
                        rowheight=24, font=FONT)
            s.configure("Treeview.Heading",
                        background=PRIMARY_PALE, foreground=PRIMARY,
                        font=FONT_BOLD, relief="flat")
            s.map("Treeview",
                  background=[("selected", PRIMARY_PALE)],
                  foreground=[("selected", PRIMARY_DARK)])
        except Exception:
            pass  # Theme failures are non-fatal

    # ── Root ───────────────────────────────────────────────────────────────────

    def _build_ui(self):
        hdr = tk.Frame(self, bg=PRIMARY_DARK, pady=8)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Hilton High School  -  Report Card Splitter",
                 bg=PRIMARY_DARK, fg="white",
                 font=("Segoe UI", 12, "bold")).pack(side="left", padx=16)

        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)

        self._build_tab_generate()
        self._build_tab_paycodes()
        self._build_tab_defaulters()
        self._build_tab_upload()
        self.bind_all("<Shift-F5>", self._secret_keypress)

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 — GENERATE
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_generate(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  Generate  ")
        outer = tk.Frame(tab, bg=BG)
        outer.pack(fill="both", expand=True, padx=10, pady=10)
        self._gen_left(outer)
        self._gen_middle(outer)
        self._gen_right(outer)

    def _gen_left(self, parent):
        col = self._card(parent, "  Source Files  ", width=230)
        col.pack(side="left", fill="y", padx=(0, 6))
        self._b(col, "+ Add File(s)", self._add_file, bg=PRIMARY, fg="white").pack(fill="x", pady=(0, 4))
        self._b(col, "Remove Selected", self._remove_file, bg=RED, fg="white").pack(fill="x", pady=(0, 10))
        tk.Label(col, text="Loaded files:", bg=CARD, font=FONT, fg=TEXT_MUTED).pack(anchor="w")
        frm = tk.Frame(col, bg=CARD); frm.pack(fill="both", expand=True)
        sb  = tk.Scrollbar(frm); sb.pack(side="right", fill="y")
        self.file_list = tk.Listbox(frm, font=("Segoe UI", 8),
                                    selectmode="single", activestyle="none",
                                    bg="white", bd=1, relief="solid",
                                    highlightthickness=0, yscrollcommand=sb.set)
        self.file_list.pack(fill="both", expand=True)
        sb.config(command=self.file_list.yview)
        self.file_list.bind("<<ListboxSelect>>", self._on_file_select)

    def _gen_middle(self, parent):
        col = self._card(parent, "  Preview  -  drag on page to mark the student name area  ")
        col.pack(side="left", fill="both", expand=True, padx=6)

        nav = tk.Frame(col, bg=CARD); nav.pack(fill="x", pady=(0, 2))
        self._b(nav, "< Prev", self._prev_page, bg=BTN_BG).pack(side="left")
        self.page_label = tk.Label(nav, text="No file loaded", bg=CARD, font=FONT, fg=TEXT_MUTED)
        self.page_label.pack(side="left", padx=10)
        self._b(nav, "Next >", self._next_page, bg=BTN_BG).pack(side="left")
        self.region_label = tk.Label(nav, text="", bg=CARD, font=FONT, fg=PRIMARY)
        self.region_label.pack(side="right", padx=6)

        tog = tk.Frame(col, bg="#dde8f0", pady=5); tog.pack(fill="x", pady=(0, 4))
        tk.Checkbutton(tog, text="Same name region for all files",
                       variable=self.same_region_var,
                       command=self._on_same_region_toggle,
                       bg="#dde8f0", font=FONT, activebackground="#dde8f0",
                       cursor="hand2").pack(side="left", padx=8)
        self.region_status_lbl = tk.Label(tog, text="", bg="#dde8f0",
                                          font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED)
        self.region_status_lbl.pack(side="right", padx=8)

        cf  = tk.Frame(col, bg="#c8d4e0"); cf.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(cf, orient="vertical")
        hsb = ttk.Scrollbar(cf, orient="horizontal")
        self.canvas = tk.Canvas(cf, bg="#c8d4e0", cursor="crosshair",
                                highlightthickness=0,
                                yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self.canvas.yview)
        hsb.config(command=self.canvas.xview)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.canvas.pack(fill="both", expand=True)
        self.canvas.bind("<ButtonPress-1>",   self._sel_start)
        self.canvas.bind("<B1-Motion>",       self._sel_drag)
        self.canvas.bind("<ButtonRelease-1>", self._sel_end)

        self.hint_lbl = tk.Label(col,
            text="Load a file, then drag on the preview to mark the student name area.",
            bg=CARD, font=("Segoe UI", 9, "italic"), fg=TEXT_MUTED)
        self.hint_lbl.pack(pady=4)

    def _gen_right(self, parent):
        col = self._card(parent, "  Settings  ", width=304)
        col.pack(side="right", fill="y", padx=(6, 0))

        self._lbl(col, "Report Card Title (for selected file):")
        self.report_name_var = tk.StringVar(value="")
        self.report_name_var.trace_add("write", self._on_report_name_changed)
        tk.Entry(col, textvariable=self.report_name_var,
                 font=("Segoe UI", 10), bd=1, relief="solid",
                 bg="white").pack(fill="x", pady=(2, 2))
        self.report_name_hint = tk.Label(col,
            text="Select a file on the left to set its title.",
            bg=CARD, font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED,
            wraplength=268, justify="left")
        self.report_name_hint.pack(anchor="w", pady=(0, 12))

        self._lbl(col, "Split Mode:")
        self.split_mode = tk.StringVar(value="fixed")
        fr = tk.Frame(col, bg=CARD); fr.pack(fill="x", pady=(2, 4))
        tk.Radiobutton(fr, text="Fixed pages per report:",
                       variable=self.split_mode, value="fixed",
                       bg=CARD, font=FONT, activebackground=CARD).pack(side="left")
        self.pages_spin = tk.Spinbox(fr, from_=1, to=100, width=5,
                                     font=("Segoe UI", 10), bd=1, relief="solid")
        self.pages_spin.delete(0, "end"); self.pages_spin.insert(0, "1")
        self.pages_spin.pack(side="left", padx=6)
        tk.Radiobutton(col, text="Auto-detect: new report when student name changes",
                       variable=self.split_mode, value="auto",
                       bg=CARD, font=FONT, justify="left",
                       activebackground=CARD, wraplength=260).pack(anchor="w", pady=(0, 14))

        self._lbl(col, "Save To (auto-creates 'report-cards' subfolder):")
        sr = tk.Frame(col, bg=CARD); sr.pack(fill="x", pady=(2, 16))
        self.save_path_var = tk.StringVar()
        tk.Entry(sr, textvariable=self.save_path_var,
                 font=FONT, bd=1, relief="solid", bg="white").pack(side="left", fill="x", expand=True)
        self._b(sr, "...", self._browse_save, bg=BTN_BG, padx=6, pady=4).pack(side="left", padx=(4, 0))

        ttk.Separator(col, orient="horizontal").pack(fill="x", pady=(0, 10))
        ch = tk.Frame(col, bg=CARD); ch.pack(fill="x")
        self._lbl(ch, "Append Circular")
        tk.Label(ch, text="optional", bg=CARD,
                 font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED).pack(side="left", padx=(6, 0))
        tk.Label(col, text="These pages are added after each student card.",
                 bg=CARD, font=FONT_SM, fg=TEXT_MUTED).pack(anchor="w", pady=(2, 6))
        self._b(col, "Choose Circular PDF", self._browse_circular,
                bg="#5b4fcf", fg="white").pack(fill="x", pady=(0, 4))
        self.circular_label = tk.Label(col, text="No circular selected",
                                       bg=CARD, font=("Segoe UI", 8, "italic"),
                                       fg=TEXT_MUTED, wraplength=250, justify="left")
        self.circular_label.pack(anchor="w")
        self._b(col, "Remove Circular", self._clear_circular, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        ttk.Separator(col, orient="horizontal").pack(fill="x", pady=(0, 10))
        self._lbl(col, "Optional Medical Form")
        self._b(col, "Choose Medical Form PDF", self._browse_medical,
                bg="#0f766e", fg="white").pack(fill="x", pady=(0, 4))
        self.medical_label = tk.Label(col, text="No medical form selected",
                                       bg=CARD, font=("Segoe UI", 8, "italic"),
                                       fg=TEXT_MUTED, wraplength=250, justify="left")
        self.medical_label.pack(anchor="w")
        self._b(col, "Remove Medical Form", self._clear_medical, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        self._lbl(col, "Optional Website Banner")
        self._b(col, "Choose Banner Image", self._browse_banner,
                bg="#0369a1", fg="white").pack(fill="x", pady=(0, 4))
        self.banner_label = tk.Label(col, text="No banner selected",
                                     bg=CARD, font=("Segoe UI", 8, "italic"),
                                     fg=TEXT_MUTED, wraplength=250, justify="left")
        self.banner_label.pack(anchor="w")
        self._b(col, "Remove Banner", self._clear_banner, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        ttk.Separator(col, orient="horizontal").pack(fill="x", pady=(0, 12))
        self.start_btn = self._b(col, "START", self._start_processing,
                                 bg=PRIMARY, fg="white", font=FONT_LG, pady=12)
        self.start_btn.pack(fill="x", pady=(0, 12))

        self._lbl(col, "Progress:")
        self.progress_var = tk.DoubleVar()
        self.pct_label = tk.Label(col, text="0%", bg=CARD, font=FONT, fg=TEXT_MUTED)
        self.pct_label.pack(anchor="e")
        ttk.Progressbar(col, variable=self.progress_var, maximum=100).pack(fill="x", pady=(0, 6))
        self.status_label = tk.Label(col, text="Ready.", bg=CARD, font=FONT_SM,
                                     fg=TEXT_MUTED, wraplength=268, justify="left")
        self.status_label.pack(anchor="w")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 — PAY CODES
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_paycodes(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  Pay Codes  ")

        tb = tk.Frame(tab, bg=PRIMARY_DARK, pady=10, padx=12); tb.pack(fill="x")
        tk.Label(tb, text="School Pay Codes",
                 bg=PRIMARY_DARK, fg="white", font=FONT_LG).pack(side="left")
        self._b(tb, "Clear All", self._pc_clear_all, bg="#7f1d1d", fg="white").pack(side="right", padx=(4, 0))
        self._b(tb, "Remove Selected", self._pc_remove_selected, bg=RED, fg="white").pack(side="right", padx=4)
        self._b(tb, "Import File", self._pc_import, bg=ACCENT, fg="white").pack(side="right")

        meta = tk.Frame(tab, bg=PRIMARY_PALE, pady=6, padx=12); meta.pack(fill="x")
        self.pc_stats = tk.Label(meta, text="", bg=PRIMARY_PALE, font=FONT_SM, fg=PRIMARY)
        self.pc_stats.pack(side="left")
        tk.Label(meta, text="Search:", bg=PRIMARY_PALE, font=FONT).pack(side="right", padx=(0, 6))
        self.pc_search_var = tk.StringVar()
        self.pc_search_var.trace_add("write", lambda *_: self._pc_refresh())
        tk.Entry(meta, textvariable=self.pc_search_var, font=FONT,
                 bd=1, relief="solid", bg="white", width=28).pack(side="right", padx=(0, 6))

        tbl = tk.Frame(tab, bg=BG); tbl.pack(fill="both", expand=True, padx=12, pady=10)
        cols = ("student_name", "codes")
        self.pc_tree = ttk.Treeview(tbl, columns=cols, show="headings", selectmode="extended")
        self.pc_tree.heading("student_name", text="Student Name")
        self.pc_tree.heading("codes",        text="School Pay Code(s)")
        self.pc_tree.column("student_name", width=300, minwidth=160)
        self.pc_tree.column("codes",        width=300, minwidth=120)
        vsb = ttk.Scrollbar(tbl, orient="vertical", command=self.pc_tree.yview)
        self.pc_tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self.pc_tree.pack(fill="both", expand=True)
        self._pc_refresh()

    def _pc_refresh(self):
        query = normalize_name(self.pc_search_var.get()) if hasattr(self, "pc_search_var") else ""
        self.pc_tree.delete(*self.pc_tree.get_children())
        students = self._store.all_students()
        shown = 0
        for name, codes in students:
            if query and query not in normalize_name(name):
                continue
            self.pc_tree.insert("", "end", values=(name, "  |  ".join(codes)))
            shown += 1
        n = len(students)
        if hasattr(self, "pc_stats"):
            self.pc_stats.config(text=f"{n} student{'s' if n != 1 else ''}  •  showing {shown}")

    def _pc_import(self):
        paths = filedialog.askopenfilenames(
            title="Import school pay codes",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("CSV", "*.csv"), ("Excel", "*.xlsx *.xls"),
                       ("Word", "*.docx *.doc"), ("PDF", "*.pdf"), ("All", "*.*")])
        if not paths:
            return
        total, errors = 0, []
        for path in paths:
            try:
                pairs = parse_paycode_file(path)
                total += self._store.import_pairs(pairs)
            except Exception as exc:
                errors.append(str(exc))
        self._pc_refresh()
        msg = f"Done. {total} new code(s) added."
        if errors:
            msg += "\n\nErrors:\n" + "\n".join(errors)
        messagebox.showinfo("Import complete", msg)

    def _pc_remove_selected(self):
        sel = self.pc_tree.selection()
        if not sel:
            messagebox.showinfo("Nothing selected", "Select one or more rows first.")
            return
        names = [self.pc_tree.item(i)["values"][0] for i in sel]
        if not messagebox.askyesno("Remove", f"Remove {len(names)} student(s)?"):
            return
        for name in names:
            self._store.remove_student(name)
        self._pc_refresh()

    def _pc_clear_all(self):
        if not messagebox.askyesno("Clear all", "Delete ALL stored pay codes?\nThis cannot be undone."):
            return
        self._store.clear_all_codes()
        self._pc_refresh()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 3 — DEFAULTERS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_defaulters(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  Defaulters  ")

        tb = tk.Frame(tab, bg="#7f1d1d", pady=10, padx=12); tb.pack(fill="x")
        tk.Label(tb, text="Fee Defaulters  -  these students will NOT receive report cards",
                 bg="#7f1d1d", fg="white", font=FONT_BOLD).pack(side="left")
        self._b(tb, "Remove Selected", self._df_remove_selected, bg="#450a0a", fg="white").pack(side="right", padx=(4, 0))
        self._b(tb, "Import Names", self._df_import, bg=AMBER, fg="white").pack(side="right")

        self.df_stats = tk.Label(tab, bg=RED_PALE, font=FONT_SM, fg=RED, pady=5, padx=12, anchor="w")
        self.df_stats.pack(fill="x")

        lst = tk.Frame(tab, bg=BG); lst.pack(fill="both", expand=True, padx=12, pady=(10, 4))
        sb  = tk.Scrollbar(lst); sb.pack(side="right", fill="y")
        self.df_listbox = tk.Listbox(lst, font=FONT, bg=CARD, bd=1, relief="solid",
                                     selectmode="extended", highlightthickness=0,
                                     yscrollcommand=sb.set)
        self.df_listbox.pack(fill="both", expand=True)
        sb.config(command=self.df_listbox.yview)

        ar = tk.Frame(tab, bg=BG, padx=12, pady=8); ar.pack(fill="x")
        tk.Label(ar, text="Add name manually:", bg=BG, font=FONT).pack(side="left")
        self.df_manual_var = tk.StringVar()
        tk.Entry(ar, textvariable=self.df_manual_var, font=FONT,
                 bd=1, relief="solid", bg="white", width=32).pack(side="left", padx=(8, 4))
        self._b(ar, "Add", self._df_add_manual, bg=RED, fg="white").pack(side="left")
        self._df_refresh()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 4 — WEBSITE UPLOAD AND SAVED REPORTS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_upload(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  Website Upload  ")
        self.upload_tab = tab

        gate = tk.Frame(tab, bg=PRIMARY_DARK, pady=10, padx=12)
        gate.pack(fill="x")
        tk.Label(gate, text="Website upload password:",
                 bg=PRIMARY_DARK, fg="white", font=FONT_BOLD).pack(side="left")
        self.upload_password_var = tk.StringVar()
        tk.Entry(gate, textvariable=self.upload_password_var, show="•",
                 font=FONT, width=24).pack(side="left", padx=8)
        self._b(gate, "Unlock Upload", self._unlock_upload,
                bg=PRIMARY, fg="white").pack(side="left")
        self.upload_gate_label = tk.Label(gate, text="Locked",
                                          bg=PRIMARY_DARK, fg="#fbbf24", font=FONT_SM)
        self.upload_gate_label.pack(side="left", padx=10)

        self.upload_body = tk.Frame(tab, bg=BG, padx=12, pady=12)
        self.upload_body.pack(fill="both", expand=True)
        self._build_history_table(self.upload_body)

    def _build_history_table(self, parent):
        toolbar = tk.Frame(parent, bg=BG)
        toolbar.pack(fill="x", pady=(0, 8))
        self._b(toolbar, "Compile Website Folder", self._compile_site,
                bg=ACCENT, fg="white").pack(side="left")
        self._b(toolbar, "Upload Folder to Netlify", self._upload_site,
                bg=PRIMARY, fg="white").pack(side="left", padx=6)
        self._b(toolbar, "Delete Selected Batch", self._delete_selected_batch,
                bg=RED, fg="white").pack(side="left", padx=6)
        self._b(toolbar, "Refresh", self._refresh_history,
                bg=BTN_BG).pack(side="left")
        self.history_storage_label = tk.Label(toolbar, text="",
                                              bg=BG, fg=TEXT_MUTED, font=FONT_SM)
        self.history_storage_label.pack(side="right")
        self.upload_status_label = tk.Label(parent, text="No website folder compiled yet.",
                                             bg=BG, fg=TEXT_MUTED, font=FONT_SM,
                                             anchor="w")
        self.upload_status_label.pack(fill="x", pady=(0, 8))

        frame = tk.Frame(parent, bg=BG)
        frame.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(frame, orient="vertical")
        vsb.pack(side="right", fill="y")
        self.history_tree = ttk.Treeview(
            frame, columns=("serial", "title", "date", "size"),
            show="headings", selectmode="extended",
            yscrollcommand=vsb.set)
        vsb.config(command=self.history_tree.yview)
        headings = (("serial", "No.", 60), ("title", "Title of reports", 350),
                    ("date", "Date", 190), ("size", "Size of all reports", 160))
        for key, text, width in headings:
            self.history_tree.heading(key, text=text)
            self.history_tree.column(key, width=width, anchor="w")
        self.history_tree.pack(fill="both", expand=True)
        self._refresh_history()

    def _refresh_history(self):
        if not hasattr(self, "history_tree"):
            return
        self.history_tree.delete(*self.history_tree.get_children())
        for number, record in enumerate(self._store.report_history, 1):
            self.history_tree.insert("", "end", iid=str(number - 1),
                                     values=(number, record.get("title", ""),
                                             record.get("date", "").replace("T", " "),
                                             DataStore.human_size(int(record.get("size", 0)))))
        used = self._store.storage_bytes()
        colour = RED if used >= MAX_REPORT_STORAGE else (
            AMBER if used >= STORAGE_WARNING_BYTES else PRIMARY)
        self.history_storage_label.config(
            text=f"Saved report storage: {DataStore.human_size(used)} / 5.0 GB",
            fg=colour)

    def _unlock_upload(self):
        if self.upload_password_var.get() == self._store.settings.get("upload_password"):
            self._upload_unlocked = True
            self.upload_gate_label.config(text="Unlocked", fg="#86efac")
            self._refresh_history()
        else:
            self._upload_unlocked = False
            self.upload_gate_label.config(text="Incorrect password", fg="#fca5a5")
            messagebox.showerror("Upload locked", "The upload password is incorrect.")

    def _delete_selected_batch(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        selected = self.history_tree.selection()
        if not selected:
            messagebox.showinfo("Nothing selected", "Select one or more saved report batches.")
            return
        records = [self._store.report_history[int(item)] for item in selected]
        if not messagebox.askyesno("Delete saved reports",
                                   f"Delete {len(records)} saved report batch(es)?"):
            return
        for record in records:
            self._store.delete_report_batch(record)
        self._refresh_history()

    def _secret_keypress(self, _event=None):
        now = time.monotonic()
        self._secret_key_presses = [stamp for stamp in self._secret_key_presses
                                    if now - stamp < 5]
        self._secret_key_presses.append(now)
        if len(self._secret_key_presses) >= 5:
            self._secret_key_presses.clear()
            self._open_secret_settings()

    def _open_secret_settings(self):
        if self._settings_tab is None:
            tab = tk.Frame(self.nb, bg=BG)
            self.nb.add(tab, text="  Secret Settings  ")
            self._settings_tab = tab
            self._build_secret_settings(tab)
        self.nb.select(self._settings_tab)

    def _build_secret_settings(self, tab):
        panel = self._card(tab, "  Administrator Settings  ")
        panel.pack(fill="x", padx=20, pady=20)
        tk.Label(panel, text="These settings are stored locally on this computer.",
                 bg=CARD, fg=TEXT_MUTED, font=FONT_SM).pack(anchor="w", pady=(0, 12))
        self.secret_vars = {
            "upload_password": tk.StringVar(value=self._store.settings.get("upload_password", "")),
            "netlify_token": tk.StringVar(value=self._store.settings.get("netlify_token", "")),
            "netlify_site_id": tk.StringVar(value=self._store.settings.get("netlify_site_id", "")),
            "copyright_year": tk.StringVar(value=self._store.settings.get("copyright_year", "")),
        }
        fields = [
            ("Upload password", "upload_password", True),
            ("Netlify personal access token", "netlify_token", True),
            ("Netlify site ID", "netlify_site_id", False),
            ("Website copyright year", "copyright_year", False),
        ]
        for label, key, masked in fields:
            row = tk.Frame(panel, bg=CARD)
            row.pack(fill="x", pady=4)
            tk.Label(row, text=label, width=30, anchor="w",
                     bg=CARD, fg=TEXT, font=FONT).pack(side="left")
            tk.Entry(row, textvariable=self.secret_vars[key],
                     show="•" if masked else "", font=FONT,
                     width=48).pack(side="left", fill="x", expand=True)
        actions = tk.Frame(panel, bg=CARD)
        actions.pack(fill="x", pady=(14, 0))
        self._b(actions, "Save Settings", self._save_secret_settings,
                bg=PRIMARY, fg="white").pack(side="left")
        self._b(actions, "Export All App Data", self._export_app_data,
                bg=ACCENT, fg="white").pack(side="left", padx=6)
        self._b(actions, "Import App Data ZIP", self._import_app_data,
                bg=BTN_BG).pack(side="left")

    def _save_secret_settings(self):
        self._store.save_settings({key: var.get() for key, var in self.secret_vars.items()})
        self.upload_password_var.set("")
        messagebox.showinfo("Saved", "Administrator settings saved.")

    def _export_app_data(self):
        path = filedialog.asksaveasfilename(
            title="Export all app data", defaultextension=".zip",
            filetypes=[("ZIP archive", "*.zip")])
        if not path:
            return
        try:
            self._store.export_archive(path)
            messagebox.showinfo("Export complete", f"All app data was exported to:\n{path}")
        except Exception as exc:
            messagebox.showerror("Export failed", str(exc))

    def _import_app_data(self):
        path = filedialog.askopenfilename(
            title="Import app data ZIP", filetypes=[("ZIP archive", "*.zip")])
        if not path:
            return
        if not messagebox.askyesno("Import app data",
                                   "Importing replaces the local saved app data. Continue?"):
            return
        try:
            self._store.import_archive(path)
            self.upload_password_var.set("")
            self._upload_unlocked = False
            for key, var in self.secret_vars.items():
                var.set(self._store.settings.get(key, ""))
            self._refresh_history()
            messagebox.showinfo("Import complete",
                                "App data imported. Click Save Settings to confirm settings.")
        except Exception as exc:
            messagebox.showerror("Import failed", str(exc))

    def _template_root(self) -> Path:
        candidates = [
            Path(getattr(sys, "_MEIPASS", Path(__file__).parent)) / "website_template",
            Path(__file__).parent / "website_template",
        ]
        for candidate in candidates:
            if (candidate / "index.html").exists():
                return candidate
        raise RuntimeError(
            "The website template is missing. Keep the website_template folder "
            "beside the application executable.")

    def _compile_site(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        destination_base = filedialog.askdirectory(
            title="Choose where to create the website folder")
        if not destination_base:
            return
        try:
            template = self._template_root()
            folder_name = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            site_dir = Path(destination_base) / folder_name
            site_dir.mkdir(parents=True, exist_ok=False)
            for name in ("index.html", "styles.css", "script.js",
                         "badge.png", "google1359a90e73c2fd7b.html"):
                source = template / name
                if source.exists():
                    shutil.copy2(source, site_dir / name)
            (site_dir / "report-cards").mkdir()

            # Copy every saved batch, newest first, and rewrite the manifest
            # filenames if two batches contain the same PDF name.
            manifest = []
            used_names = set()
            for batch_index, record in enumerate(self._store.report_history):
                folder = Path(record.get("folder", ""))
                rows_by_file = {
                    str(row[0]): row for row in record.get("manifest_rows", []) if row}
                for original in record.get("files", []):
                    source = folder / original
                    if not source.exists():
                        continue
                    target_name = original
                    if target_name in used_names:
                        target_name = f"{batch_index + 1}_{original}"
                    used_names.add(target_name)
                    shutil.copy2(source, site_dir / "report-cards" / target_name)
                    row = rows_by_file.get(original)
                    if row:
                        manifest.append((target_name, row[1], row[2]))
                    else:
                        stem = Path(original).stem.split("_")
                        student = stem[0] if stem else "Unknown"
                        manifest.append((target_name, student, ""))

            with open(site_dir / "report-cards" / "manifest.csv", "w",
                      newline="", encoding="utf-8") as handle:
                writer = csv.writer(handle)
                writer.writerow(["filename", "student_name", "school_pay_code"])
                writer.writerows(manifest)

            assets = site_dir / "assets"
            assets.mkdir()
            if self.circular_path and Path(self.circular_path).exists():
                shutil.copy2(self.circular_path, assets / "circular.pdf")
            if self.medical_path and Path(self.medical_path).exists():
                shutil.copy2(self.medical_path, assets / "medical-form.pdf")
            if self.banner_path and Path(self.banner_path).exists():
                ext = Path(self.banner_path).suffix.lower()
                if ext not in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
                    raise RuntimeError("The banner must be a supported image file.")
                shutil.copy2(self.banner_path, assets / f"banner{ext}")

            script = site_dir / "script.js"
            year = self._store.settings.get("copyright_year", str(datetime.now().year))
            if script.exists():
                content = script.read_text("utf-8")
                content = content.replace(
                    "document.getElementById('footer-year').textContent = new Date().getFullYear();",
                    "document.getElementById('footer-year').textContent = " +
                    json.dumps(str(year)) + ";")
                script.write_text(content, "utf-8")

            self.compiled_site_dir = site_dir
            self.upload_status_label.config(
                text=f"Compiled {len(manifest)} report(s) into {site_dir}",
                fg=PRIMARY)
            messagebox.showinfo("Website compiled",
                                f"Website folder created:\n{site_dir}\n\n"
                                f"{len(manifest)} report(s) included.")
        except Exception as exc:
            messagebox.showerror("Compile failed", str(exc))

    def _zip_folder(self, folder: Path, destination: Path):
        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in folder.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(folder))

    def _upload_site(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        folder = getattr(self, "compiled_site_dir", None)
        if not folder or not Path(folder).exists():
            messagebox.showinfo("Compile first",
                                "Compile a website folder before uploading it.")
            return
        token = self._store.settings.get("netlify_token", "").strip()
        site_id = self._store.settings.get("netlify_site_id", "").strip()
        if not token or not site_id:
            messagebox.showerror("Missing Netlify settings",
                                 "Set the Netlify token and site ID in Secret Settings.")
            return
        self.upload_status_label.config(text="Preparing upload...", fg=ACCENT)
        threading.Thread(target=self._upload_site_worker,
                         args=(Path(folder), token, site_id), daemon=True).start()

    def _upload_site_worker(self, folder: Path, token: str, site_id: str):
        archive_path = None
        try:
            fd, temp_name = tempfile.mkstemp(suffix=".zip")
            os.close(fd)
            archive_path = Path(temp_name)
            self.after(0, lambda: self.upload_status_label.config(
                text="Compressing website folder...", fg=ACCENT))
            self._zip_folder(folder, archive_path)
            payload = archive_path.read_bytes()
            self.after(0, lambda: self.upload_status_label.config(
                text=f"Uploading {DataStore.human_size(len(payload))} to Netlify...",
                fg=ACCENT))
            request = urllib.request.Request(
                f"https://api.netlify.com/api/v1/sites/{site_id}/deploys",
                data=payload,
                method="POST",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/zip",
                    "Content-Length": str(len(payload)),
                })
            with urllib.request.urlopen(request, timeout=180) as response:
                response.read()
            self.after(0, lambda: self.upload_status_label.config(
                text="Upload complete. The website is deploying on Netlify.",
                fg=PRIMARY))
            self.after(0, lambda: messagebox.showinfo(
                "Upload complete",
                "The complete website folder was uploaded to Netlify successfully."))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")[:500]
            code = exc.code
            self.after(0, lambda: messagebox.showerror(
                "Netlify upload failed", f"Netlify returned {code}.\n{detail}"))
            self.after(0, lambda: self.upload_status_label.config(
                text="Upload failed. Check Netlify settings and try again.", fg=RED))
        except Exception as exc:
            error_text = str(exc)
            self.after(0, lambda: messagebox.showerror("Upload failed", error_text))
            self.after(0, lambda: self.upload_status_label.config(
                text=f"Upload failed: {error_text}", fg=RED))
        finally:
            if archive_path:
                try:
                    archive_path.unlink()
                except OSError:
                    pass

    def _df_refresh(self):
        self.df_listbox.delete(0, "end")
        for name in self._store.all_defaulters():
            self.df_listbox.insert("end", name)
        n = len(self._store.all_defaulters())
        self.df_stats.config(text=f"  {n} student{'s' if n != 1 else ''} on the defaulters list")

    def _df_import(self):
        paths = filedialog.askopenfilenames(
            title="Import defaulter names",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("All", "*.*")])
        if not paths:
            return
        total, errors = 0, []
        for path in paths:
            try:
                names = parse_names_file(path)
                self._store.import_defaulters(names)
                total += len(names)
            except Exception as exc:
                errors.append(str(exc))
        self._df_refresh()
        msg = f"Done. {total} name(s) added to defaulters list."
        if errors:
            msg += "\n\nErrors:\n" + "\n".join(errors)
        messagebox.showinfo("Import complete", msg)

    def _df_remove_selected(self):
        sel = list(self.df_listbox.curselection())
        if not sel:
            messagebox.showinfo("Nothing selected", "Select one or more names first.")
            return
        names = [self.df_listbox.get(i) for i in sel]
        if not messagebox.askyesno("Remove", f"Remove {len(names)} name(s)?"):
            return
        for name in names:
            self._store.remove_defaulter(name)
        self._df_refresh()

    def _df_add_manual(self):
        name = self.df_manual_var.get().strip()
        if name:
            self._store.add_defaulter(name)
            self.df_manual_var.set("")
            self._df_refresh()

    # WIDGET HELPERS
    # ══════════════════════════════════════════════════════════════════════════

    def _card(self, parent, title="", **kw):
        return tk.LabelFrame(parent, text=title, bg=CARD,
                             font=FONT_BOLD, padx=10, pady=10,
                             bd=1, relief="solid", **kw)

    def _b(self, parent, text, command, bg=BTN_BG, fg=TEXT,
           font=None, padx=10, pady=5):
        return tk.Button(parent, text=text, command=command,
                         bg=bg, fg=fg, font=font or FONT,
                         relief="flat", padx=padx, pady=pady,
                         cursor="hand2", activebackground=bg, activeforeground=fg)

    def _lbl(self, parent, text):
        tk.Label(parent, text=text, bg=CARD,
                 font=FONT_BOLD, fg=TEXT).pack(anchor="w")

    # ══════════════════════════════════════════════════════════════════════════
    # REGION HELPERS
    # ══════════════════════════════════════════════════════════════════════════

    def _cur_rect(self):
        return self.name_rects.get(self.current_path)

    def _on_same_region_toggle(self):
        if self.same_region_var.get():
            rect = self._cur_rect()
            if rect:
                for p in self.file_list.get(0, "end"):
                    self.name_rects[p] = rect
        self._update_region_status()

    def _update_region_status(self):
        files = list(self.file_list.get(0, "end"))
        n = len(files)
        if n == 0:
            self.region_status_lbl.config(text="")
            return
        if self.same_region_var.get():
            if self._cur_rect():
                self.region_status_lbl.config(text="Region set - applies to all files", fg=PRIMARY)
            else:
                self.region_status_lbl.config(text="No region set yet - drag on the preview", fg=TEXT_MUTED)
        else:
            have  = sum(1 for p in files if p in self.name_rects)
            color = PRIMARY if have == n else (RED if have == 0 else AMBER)
            self.region_status_lbl.config(
                text=f"{have} of {n} file{'s' if n != 1 else ''} have a region set", fg=color)

    # ══════════════════════════════════════════════════════════════════════════
    # FILE MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════

    def _add_file(self):
        paths = filedialog.askopenfilenames(
            title="Select report card PDF(s)",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        existing = list(self.file_list.get(0, "end"))
        added = False
        for p in paths:
            if p not in existing:
                self.file_list.insert("end", p)
                if self.same_region_var.get() and self.name_rects:
                    self.name_rects[p] = next(iter(self.name_rects.values()))
                # Initialise per-file title (inherit current entry value if set)
                if p not in self.file_report_names:
                    self.file_report_names[p] = self.report_name_var.get().strip()
                added = True
        if added:
            self._update_region_status()
            if self.pdf_doc is None:
                self.file_list.selection_set(0)
                self._load_selected()

    def _remove_file(self):
        sel = self.file_list.curselection()
        if not sel:
            return
        path = self.file_list.get(sel[0])
        self.file_list.delete(sel[0])
        self.name_rects.pop(path, None)
        self.file_report_names.pop(path, None)
        if path == self.current_path:
            self.pdf_doc = None; self.current_path = ""
            self.canvas.delete("all")
            self.page_label.config(text="No file loaded")
            self.region_label.config(text="")
            self.hint_lbl.config(
                text="Load a file, then drag on the preview to mark the student name area.",
                fg=TEXT_MUTED)
        self._update_region_status()

    def _on_file_select(self, _=None):
        self._load_selected()

    def _on_report_name_changed(self, *_):
        """Save the current entry value back to the per-file dict."""
        if self.current_path:
            self.file_report_names[self.current_path] = self.report_name_var.get()

    def _load_selected(self):
        sel = self.file_list.curselection()
        if sel:
            self._load_file(self.file_list.get(sel[0]))

    def _load_file(self, path):
        try:
            doc = fitz.open(path)
        except Exception as exc:
            messagebox.showerror("Error", f"Could not open file:\n{exc}"); return
        # Set current_path BEFORE updating the StringVar so the trace writes correctly
        self.pdf_doc, self.current_path = doc, path
        self.current_page = 0; self.sel_rect_id = None
        # Load this file's individual report card title into the entry
        self.report_name_var.set(self.file_report_names.get(path, ""))
        fname = Path(path).name
        self.report_name_hint.config(text=f"Title for: {fname}")
        if self._cur_rect():
            self.region_label.config(text="Region already set for this file", fg=PRIMARY)
            self.hint_lbl.config(
                text="Name region is marked (shown in orange). Drag again to change it.",
                fg=TEXT_MUTED)
        else:
            self.region_label.config(text="")
            self.hint_lbl.config(
                text="Drag on the preview to mark where the student name appears.",
                fg=TEXT_MUTED)
        self._render_page()
        self._update_region_status()

    # ══════════════════════════════════════════════════════════════════════════
    # PDF RENDERING
    # ══════════════════════════════════════════════════════════════════════════

    def _fit_canvas(self):
        if self.pdf_doc:
            self._render_page()

    def _render_page(self):
        if not self.pdf_doc:
            return
        page  = self.pdf_doc[self.current_page]
        total = len(self.pdf_doc)
        self.page_label.config(text=f"Page {self.current_page + 1} of {total}")
        self.canvas.update_idletasks()
        cw = max(self.canvas.winfo_width(), 400)
        self.page_scale = cw / page.rect.width
        mat = fitz.Matrix(self.page_scale, self.page_scale)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        self.page_photo = ImageTk.PhotoImage(img)
        self.canvas.delete("all")
        self.canvas.create_image(0, 0, anchor="nw", image=self.page_photo)
        self.canvas.config(scrollregion=(0, 0, pix.width, pix.height))
        self.sel_rect_id = None
        rect = self._cur_rect()
        if rect:
            self._draw_region(rect)

    def _draw_region(self, rect):
        s = self.page_scale
        self.sel_rect_id = self.canvas.create_rectangle(
            rect.x0*s, rect.y0*s, rect.x1*s, rect.y1*s,
            outline="#f59e0b", width=2, dash=(5, 3))

    def _prev_page(self):
        if self.pdf_doc and self.current_page > 0:
            self.current_page -= 1; self._render_page()

    def _next_page(self):
        if self.pdf_doc and self.current_page < len(self.pdf_doc) - 1:
            self.current_page += 1; self._render_page()

    # ══════════════════════════════════════════════════════════════════════════
    # REGION SELECTION
    # ══════════════════════════════════════════════════════════════════════════

    def _sel_start(self, event):
        self.sel_start = (self.canvas.canvasx(event.x), self.canvas.canvasy(event.y))
        if self.sel_rect_id:
            self.canvas.delete(self.sel_rect_id); self.sel_rect_id = None

    def _sel_drag(self, event):
        if not self.sel_start:
            return
        x0, y0 = self.sel_start
        x1, y1 = self.canvas.canvasx(event.x), self.canvas.canvasy(event.y)
        if self.sel_rect_id:
            self.canvas.coords(self.sel_rect_id, x0, y0, x1, y1)
        else:
            self.sel_rect_id = self.canvas.create_rectangle(
                x0, y0, x1, y1, outline="#f59e0b", width=2, dash=(5, 3))

    def _sel_end(self, event):
        if not self.sel_start:
            return
        x0, y0 = self.sel_start
        x1, y1 = self.canvas.canvasx(event.x), self.canvas.canvasy(event.y)
        self.sel_start = None
        rx0, rx1 = sorted([x0, x1]); ry0, ry1 = sorted([y0, y1])
        if (rx1 - rx0) < 8 or (ry1 - ry0) < 8:
            return
        s    = self.page_scale
        rect = fitz.Rect(rx0/s, ry0/s, rx1/s, ry1/s)
        if self.same_region_var.get():
            for p in self.file_list.get(0, "end"):
                self.name_rects[p] = rect
            scope = "applied to all files"
        else:
            self.name_rects[self.current_path] = rect
            scope = "set for this file"
        self.region_label.config(
            text=f"Region {scope}  ({int(rx0/s)},{int(ry0/s)}) > ({int(rx1/s)},{int(ry1/s)})",
            fg=PRIMARY)
        sample = self._extract_name(self.pdf_doc, self.current_page, rect)
        if sample:
            self.hint_lbl.config(text=f'Sample name from this page: "{sample}"', fg=PRIMARY)
        else:
            self.hint_lbl.config(
                text="No text found in that area. Try a larger or different region.", fg=RED)
        self._update_region_status()

    # ══════════════════════════════════════════════════════════════════════════
    # SETTINGS
    # ══════════════════════════════════════════════════════════════════════════

    def _browse_save(self):
        path = filedialog.askdirectory(title="Choose base output folder")
        if path:
            self.save_path_var.set(path)

    def _browse_circular(self):
        path = filedialog.askopenfilename(
            title="Choose circular PDF",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        if not path:
            return
        try:
            doc   = fitz.open(path)
            pages = len(doc)
            self.circular_doc = doc
            self.circular_path = path
            self.circular_label.config(
                text=f"OK  {Path(path).name}  ({pages} page{'s' if pages != 1 else ''})",
                fg=PRIMARY)
        except Exception as exc:
            messagebox.showerror("Error", f"Could not open circular PDF:\n{exc}")

    def _clear_circular(self):
        self.circular_doc = None
        self.circular_path = ""
        self.circular_label.config(text="No circular selected", fg=TEXT_MUTED)

    def _browse_medical(self):
        path = filedialog.askopenfilename(
            title="Choose medical form PDF",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        if path:
            self.medical_path = path
            self.medical_label.config(text=f"OK  {Path(path).name}", fg=PRIMARY)

    def _clear_medical(self):
        self.medical_path = ""
        self.medical_label.config(text="No medical form selected", fg=TEXT_MUTED)

    def _browse_banner(self):
        path = filedialog.askopenfilename(
            title="Choose website banner image",
            filetypes=[("Images", "*.png *.jpg *.jpeg *.gif *.webp *.bmp"),
                       ("All", "*.*")])
        if path:
            self.banner_path = path
            self.banner_label.config(text=f"OK  {Path(path).name}", fg=PRIMARY)

    def _clear_banner(self):
        self.banner_path = ""
        self.banner_label.config(text="No banner selected", fg=TEXT_MUTED)

    # ══════════════════════════════════════════════════════════════════════════
    # PROCESSING
    # ══════════════════════════════════════════════════════════════════════════

    def _start_processing(self):
        if self.processing:
            return
        files = list(self.file_list.get(0, "end"))
        if not files:
            messagebox.showerror("No files", "Please add at least one PDF file."); return
        missing_regions = [p for p in files if p not in self.name_rects]
        if missing_regions:
            messagebox.showerror("Missing name region",
                "These files have no name region marked:\n\n" +
                "\n".join(f"  - {Path(p).name}" for p in missing_regions) +
                "\n\nClick each file and drag on the preview to mark the name area.")
            return
        # Collect per-file report card titles; validate all are filled
        file_report_names_snap = {}
        missing_titles = []
        for p in files:
            title = self.file_report_names.get(p, "").strip()
            if not title:
                missing_titles.append(Path(p).name)
            file_report_names_snap[p] = title
        if missing_titles:
            messagebox.showerror("Missing Report Card Title",
                "Please set a Report Card Title for each file.\n\n"
                "Select each file in the list on the left and enter its title.\n\n"
                "Missing titles for:\n" +
                "\n".join(f"  - {n}" for n in missing_titles))
            return
        base_path = self.save_path_var.get().strip()
        if not base_path:
            messagebox.showerror("Missing", "Please choose a Save To folder."); return
        save_path = os.path.join(base_path, "report-cards")
        os.makedirs(save_path, exist_ok=True)
        mode = self.split_mode.get()
        pages_per = 1
        if mode == "fixed":
            try:
                pages_per = int(self.pages_spin.get()); assert pages_per >= 1
            except (ValueError, AssertionError):
                messagebox.showerror("Invalid", "Pages per report must be >= 1."); return
        self.processing = True
        self.start_btn.config(state="disabled", text="Processing...")
        self._set_status(0, "Starting...")
        threading.Thread(target=self._run_processing,
                         args=(files, mode, pages_per, file_report_names_snap, save_path),
                         daemon=True).start()

    def _run_processing(self, files, mode, pages_per, file_report_names_snap, save_path):
        try:
            # Phase 1: scan all files
            all_cards = []
            for fi, file_path in enumerate(files):
                self._set_status(int(fi / len(files) * 20),
                    f"Scanning {fi+1}/{len(files)}: {Path(file_path).name}")
                src  = fitz.open(file_path)
                rect = self.name_rects[file_path]
                tot  = len(src)
                if mode == "fixed":
                    for start in range(0, tot, pages_per):
                        end  = min(start + pages_per, tot)
                        name = self._extract_name(src, start, rect) or f"Student {start//pages_per+1}"
                        all_cards.append((file_path, name, list(range(start, end))))
                else:
                    for name, group in self._auto_detect(src, rect, tot):
                        all_cards.append((file_path, name, group))
                src.close()

            # Phase 2: filter defaulters
            to_process = [(fp, nm, grp) for fp, nm, grp in all_cards
                          if not self._store.is_defaulter(nm)]
            skipped    = len(all_cards) - len(to_process)

            # Phase 3: check pay codes
            missing_codes = list({nm for _, nm, _ in to_process
                                  if not self._store.get_codes(nm)})
            if missing_codes:
                decision_holder = [None]
                done_evt = threading.Event()

                def _show():
                    dlg = MissingCodesDialog(self, missing_codes, self._store)
                    self.wait_window(dlg)
                    decision_holder[0] = dlg.result
                    done_evt.set()

                self.after(0, _show)
                done_evt.wait()
                decision = decision_holder[0]
                if decision == "cancel":
                    self._set_status(0, "Cancelled."); return
                if decision == "skip":
                    to_process = [(fp, nm, grp) for fp, nm, grp in to_process
                                  if self._store.get_codes(nm)]

            # Phase 4: generate PDFs
            manifest_rows = []
            generated_by_title = {}
            total_cards   = len(to_process)
            for i, (file_path, student_name, group) in enumerate(to_process):
                self._set_status(20 + int(i / total_cards * 80),
                    f"Saving {i+1}/{total_cards}: {student_name}")
                src      = fitz.open(file_path)
                rpt_name = file_report_names_snap.get(file_path, "Report Card")
                filename = self._save_pdf(src, group, student_name, rpt_name, save_path)
                src.close()
                codes = self._store.get_codes(student_name)
                manifest_rows.append((filename, student_name, "|".join(codes)))
                generated_by_title.setdefault(rpt_name, []).append(filename)

            # Phase 5: manifest.csv
            with open(os.path.join(save_path, "manifest.csv"), "w",
                      newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(["filename", "student_name", "school_pay_code"])
                for row in manifest_rows:
                    w.writerow(row)

            for title, filenames in generated_by_title.items():
                rows_for_title = [row for row in manifest_rows if row[0] in filenames]
                self._store.add_report_batch(
                    title, save_path, filenames, manifest_rows=rows_for_title)

            count = len(manifest_rows)
            self._set_status(100,
                f"Done! {count} report card(s) saved.\n"
                f"Defaulters skipped: {skipped}.\n"
                f"manifest.csv written to report-cards/\n\n"
                f"Upload the report-cards folder to your Netlify site.")
            self.after(0, lambda: messagebox.showinfo("Complete",
                f"{count} report card(s) saved.\n"
                f"{skipped} student(s) skipped (defaulters).\n\n"
                f"Output: {save_path}\n\n"
                "Upload the report-cards folder to your Netlify site to publish."))
            self.after(0, self._pc_refresh)
            self.after(0, self._refresh_history)

        except Exception as exc:
            self.after(0, lambda: messagebox.showerror("Error", str(exc)))
            self._set_status(0, f"Error: {exc}")
        finally:
            self.processing = False
            self.after(0, lambda: self.start_btn.config(state="normal", text="START"))

    def _auto_detect(self, source_doc, name_rect, total_pages):
        groups, cur_name, cur_pages = [], None, []
        for p in range(total_pages):
            name = self._extract_name(source_doc, p, name_rect) or cur_name or "Unknown"
            if name != cur_name:
                if cur_pages:
                    groups.append((cur_name or "Unknown", cur_pages))
                cur_name, cur_pages = name, [p]
            else:
                cur_pages.append(p)
        if cur_pages:
            groups.append((cur_name or "Unknown", cur_pages))
        return groups

    def _save_pdf(self, source_doc, page_nums, student_name, report_name, save_path):
        filename = (f"{safe_filename(student_name)}_"
                    f"{safe_filename(report_name)}_"
                    f"{random_number()}.pdf")
        new_doc = fitz.open()
        for p in page_nums:
            new_doc.insert_pdf(source_doc, from_page=p, to_page=p)
        if self.circular_doc:
            new_doc.insert_pdf(self.circular_doc)
        new_doc.save(os.path.join(save_path, filename), garbage=4, deflate=True, clean=True)
        new_doc.close()
        return filename

    @staticmethod
    def _extract_name(source_doc, page_num: int, rect) -> str:
        if not source_doc or rect is None:
            return ""
        text = source_doc[page_num].get_text("text", clip=rect).strip()
        return re.sub(r"\s+", " ", text).strip()

    def _set_status(self, pct: float, text: str):
        self.after(0, lambda: self.progress_var.set(pct))
        self.after(0, lambda: self.pct_label.config(text=f"{int(pct)}%"))
        self.after(0, lambda: self.status_label.config(text=text))


# ═══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    try:
        app = ReportSplitter()
        app.mainloop()
    except Exception as exc:
        import traceback
        try:
            root = tk.Tk(); root.withdraw()
            messagebox.showerror("Startup Error",
                f"The app failed to start:\n\n{exc}\n\n"
                f"{traceback.format_exc()[-800:]}")
            root.destroy()
        except Exception:
            pass


if __name__ == "__main__":
    main()
